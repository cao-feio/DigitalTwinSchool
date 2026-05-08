
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    return para

def add_list_item(doc, text, level=0):
    """添加列表项"""
    para = doc.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * level)
    for run in para.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
    return para

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.bold = True
    
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10.5)
    
    return table

def convert_markdown_to_word(md_file, output_file):
    """将Markdown转换为Word"""
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        if line.startswith('#'):
            level = line.count('#')
            text = line.lstrip('#').strip()
            add_heading(doc, text, min(level, 3))
            i += 1
            continue
        
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        if line.strip().startswith(('- ', '* ')):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            text = line.strip()[2:].strip()
            add_list_item(doc, text, level)
            i += 1
            continue
        
        text = line
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        add_paragraph(doc, text)
        i += 1
    
    doc.save(output_file)
    print(f'Word文档已生成: {output_file}')

if __name__ == '__main__':
    md_file = 'PRD.md'
    output_file = '数字孪生校园平台需求文档.docx'
    convert_markdown_to_word(md_file, output_file)
